"""Renders a client's Word template with ``docxtpl``.

The client's original layout — tables, headings, spacing, logos, signature
blocks — is preserved because we render *their* file; only the ``{{ tag }}``
placeholders are substituted.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docxtpl import DocxTemplate
from jinja2 import (  # noqa: F401  (documented default)
    ChainableUndefined,
    Environment,
    StrictUndefined,
)

from app.core.errors import ValidationError

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)")


def is_ooxml(path: Path) -> bool:
    """True for a real .docx (a ZIP package), false for legacy binary .doc."""
    try:
        with path.open("rb") as handle:
            return handle.read(4) == b"PK\x03\x04"
    except OSError:
        return False


def list_placeholders(path: Path) -> list[str]:
    """Every ``{{ name }}`` the template asks for, de-duplicated and sorted."""
    if not is_ooxml(path):
        return []
    template = DocxTemplate(str(path))
    try:
        names = set(template.get_undeclared_template_variables())
    except Exception:  # noqa: BLE001 - fall back to a text scan
        names = set()
    if not names:
        document = Document(str(path))
        text_blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_blocks.append(cell.text)
        for block in text_blocks:
            names.update(PLACEHOLDER_RE.findall(block))
    return sorted(names)


def render(template_path: Path, context: dict[str, Any], output_path: Path) -> Path:
    """Render ``template_path`` with ``context`` and write ``output_path``."""
    if not template_path.exists():
        raise ValidationError(
            "The document template file is missing from storage. "
            "Re-upload the template under Administration → Document Templates."
        )
    if not is_ooxml(template_path):
        raise ValidationError(
            "This template is a legacy binary .doc file. Re-save it as .docx and "
            "upload it again to enable Word generation."
        )

    document = DocxTemplate(str(template_path))
    environment = Environment(undefined=ChainableUndefined, autoescape=False)
    try:
        document.render(context, environment)
    except Exception as exc:  # noqa: BLE001 - surface template problems clearly
        raise ValidationError(
            "The document template could not be rendered.",
            details=str(exc),
        ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
