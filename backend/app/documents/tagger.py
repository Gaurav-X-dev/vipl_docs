"""Turns a *filled specimen* .docx into a reusable ``{{ tag }}`` template.

The insurer forms supplied by the client are completed real reports, not blank
templates. This module replaces the specimen values with Jinja placeholders while
leaving the surrounding layout — tables, headings, spacing, logos, signature
blocks — completely untouched.

Word splits a single visible string across several ``<w:r>`` runs (spell-check
marks, formatting changes, rsid noise), so a naive per-run replace misses most
matches. :func:`replace_in_paragraph` therefore works on the paragraph's joined
text and redistributes the result back across its runs, keeping the first run's
formatting for the replacement.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass(slots=True)
class TagReport:
    """What the tagging pass actually changed — printed by the seed script."""

    file: str
    replacements: dict[str, int] = field(default_factory=dict)
    not_found: list[str] = field(default_factory=list)

    @property
    def total(self) -> int:
        return sum(self.replacements.values())

    def as_lines(self) -> list[str]:
        lines = [f"{self.file}: {self.total} substitution(s)"]
        for needle, count in sorted(self.replacements.items(), key=lambda item: -item[1]):
            lines.append(f"    {count:>3} x {needle!r}")
        for needle in self.not_found:
            lines.append(f"      -  not found: {needle!r}")
        return lines


def _iter_paragraphs(container) -> list[Paragraph]:
    """Every paragraph in a document part, including those nested in tables."""
    paragraphs: list[Paragraph] = list(getattr(container, "paragraphs", []))
    for table in getattr(container, "tables", []):
        paragraphs.extend(_iter_table_paragraphs(table))
    return paragraphs


def _iter_table_paragraphs(table: Table) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
            for nested in cell.tables:
                paragraphs.extend(_iter_table_paragraphs(nested))
    return paragraphs


def document_paragraphs(document: Document) -> list[Paragraph]:
    paragraphs = _iter_paragraphs(document)
    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if part is not None:
                paragraphs.extend(_iter_paragraphs(part))
    return paragraphs


def replace_in_paragraph(paragraph: Paragraph, needle: str, replacement: str) -> int:
    """Replace ``needle`` with ``replacement`` across run boundaries.

    Returns the number of substitutions made in this paragraph.
    """
    if not needle or not paragraph.runs:
        return 0
    original = "".join(run.text for run in paragraph.runs)
    if needle not in original:
        return 0

    count = original.count(needle)
    updated = original.replace(needle, replacement)

    # Put the whole result in the first run and blank the rest. The first run
    # carries the paragraph's dominant formatting, which is what the reader sees.
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""
    return count


def tag_document(
    source: Path,
    destination: Path,
    replacements: dict[str, str],
    *,
    label: str | None = None,
) -> TagReport:
    """Write a tagged copy of ``source`` to ``destination``.

    ``replacements`` maps a literal specimen value to its placeholder, e.g.
    ``{"MRS. ABDA  BEGUM": "{{ life_assured_name }}"}``. Longer needles are
    applied first so that a value never gets clipped by a shorter substring.
    """
    document = Document(str(source))
    paragraphs = document_paragraphs(document)
    report = TagReport(file=label or source.name)

    for needle in sorted(replacements, key=len, reverse=True):
        placeholder = replacements[needle]
        hits = 0
        for paragraph in paragraphs:
            hits += replace_in_paragraph(paragraph, needle, placeholder)
        if hits:
            report.replacements[needle] = hits
        else:
            report.not_found.append(needle)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return report


def extract_text(source: Path) -> str:
    """Flatten a document to text — used to hunt for specimen values."""
    document = Document(str(source))
    return "\n".join(p.text for p in document_paragraphs(document))
